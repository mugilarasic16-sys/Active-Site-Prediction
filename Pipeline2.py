# =========================================================
# FINAL STABLE VERSION: 3D CAPTURE + MANUAL TRIGGER
# =========================================================

import os, subprocess, sys, base64, time
from Bio.PDB import PDBList, PDBParser
from docx import Document
from docx.shared import Inches, RGBColor
from google.colab import files, output
from IPython.display import Javascript, display, HTML
import py3Dmol

# 1. SETUP
def prepare_env():
    packages = ["biopython", "python-docx", "py3Dmol"]
    for p in packages:
        try:
            __import__(p)
        except ImportError:
            subprocess.check_call([sys.executable, "-m", "pip", "install", p, "--quiet"])
prepare_env()

# 2. INPUT
uploaded = files.upload()
file_path = list(uploaded.keys())[0] if uploaded else ""
if not file_path:
    pdb_id = input("OR Enter PDB ID: ").strip().lower()
    pdbl = PDBList()
    file_path = pdbl.retrieve_pdb_file(pdb_id, pdir='.', file_format='pdb')
    name_tag = pdb_id
else:
    name_tag = file_path.split('.')[0]

# 3. 3D VISUALIZATION & CAPTURE
def show_protein(path):
    view = py3Dmol.view(width=800, height=600)
    with open(path, 'r') as f:
        view.addModel(f.read(), 'pdb')
    view.setStyle({'cartoon': {'color': 'spectrum'}})
    view.addStyle({'resn': 'HIS'}, {'stick': {'color': 'red'}})
    view.addStyle({'resn': 'SER'}, {'stick': {'color': 'blue'}})
    view.addStyle({'resn': 'ASP'}, {'stick': {'color': 'green'}})
    view.zoomTo()
    view.show()

    # Image capture callback
    def save_image(img_data):
        with open('protein_view.png', 'wb') as f:
            f.write(base64.b64decode(img_data.split(',')[1]))
        print("\n✅ 3D Screenshot captured! You can now click the button below.")

    output.register_callback('notebook.download_image', save_image)
    display(Javascript('''
        (async function() {
            await new Promise(r => setTimeout(r, 3000));
            const canvas = document.querySelector('canvas');
            const imgData = canvas.toDataURL("image/png");
            google.colab.kernel.invokeFunction('notebook.download_image', [imgData], {});
        })();
    '''))

# 4. FINAL DOCX GENERATION (Triggered by button)
def make_docx():
    parser = PDBParser(QUIET=True)
    structure = parser.get_structure(name_tag, file_path)
    doc = Document()
    doc.add_heading(f'3D Active Site Analysis: {name_tag.upper()}', 0)

    # Insert Image
    if os.path.exists('protein_view.png'):
        doc.add_heading('Structural Visualization', level=1)
        doc.add_picture('protein_view.png', width=Inches(5.5))
        doc.add_paragraph('Figure 1: 3D protein structure (Red=HIS, Blue=SER, Green=ASP).')
    else:
        doc.add_paragraph("⚠️ IMAGE NOT FOUND. Please wait longer for the render.")

    # Table
    doc.add_heading('Activity-Based Residue Positions', level=1)
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    hdr = table.rows[0].cells
    hdr[0].text, hdr[1].text, hdr[2].text = 'HIS (Catalytic)', 'SER (Nucleophilic)', 'ASP (Stabilizer)'

    res_map = {'HIS': [], 'SER': [], 'ASP': []}
    colors = {'HIS': RGBColor(200, 0, 0), 'SER': RGBColor(0, 0, 200), 'ASP': RGBColor(0, 120, 0)}

    for model in structure:
        for chain in model:
            for res in chain:
                if res.resname in res_map and res.id[0] == ' ':
                    res_map[res.resname].append(f"{res.resname}{res.id[1]}({chain.id})")

    max_rows = max(len(res_map['HIS']), len(res_map['SER']), len(res_map['ASP']))
    for i in range(max_rows):
        row = table.add_row().cells
        for idx, key in enumerate(['HIS', 'SER', 'ASP']):
            if i < len(res_map[key]):
                run = row[idx].paragraphs[0].add_run(res_map[key][i])
                run.font.color.rgb = colors[key]

    doc.save('Final_3D_Protein_Report.docx')
    files.download('Final_3D_Protein_Report.docx')

# Register Button Callback
output.register_callback('create_docx_trigger', make_docx)

# RUN
show_protein(file_path)
display(HTML('''
    <div style="border: 2px solid #4CAF50; padding: 15px; margin-top: 20px; border-radius: 10px; background-color: #f9f9f9;">
        <p style="color: #2e7d32; font-weight: bold;">Step 2: Finalize Report</p>
        <p>Wait for the "✅ 3D Screenshot captured!" message above, then click:</p>
        <button onclick="google.colab.kernel.invokeFunction('create_docx_trigger', [], {})" 
                style="background-color: #4CAF50; color: white; padding: 10px 20px; border: none; border-radius: 5px; cursor: pointer; font-size: 16px;">
            CREATE AND DOWNLOAD DOCX
        </button>
    </div>
'''))